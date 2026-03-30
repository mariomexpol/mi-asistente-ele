import streamlit as st
import requests
import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

st.set_page_config(page_title="Asistente ELE Pro", layout="wide")

# --- LIMPIEZA DE TEXTO ---
def limpiar_texto_ele(texto):
    return texto.replace('\\_', '_').replace('\\', '')

# --- GENERADOR DE WORD PROFESIONAL ---
def generar_docx_profesional(texto_ia, escuela, profe, tema, logo_file=None):
    doc = Document()
    section = doc.sections[0]
    header = section.header
    htxt = header.paragraphs[0]
    
    if logo_file:
        try:
            run_logo = htxt.add_run()
            run_logo.add_picture(logo_file, width=Inches(1.2))
            htxt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        except: pass
    
    info_h = header.add_paragraph(f"{escuela}\nProf. {profe}")
    info_h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph() 
    titulo = doc.add_heading(tema.upper(), 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    lineas = limpiar_texto_ele(texto_ia).split('\n')
    for linea in lineas:
        l = linea.strip()
        if not l: continue
        
        # Detección de títulos para saltos de página
        if l.startswith('#') or "VOCABULARIO" in l.upper() or "SOLUCIONARIO" in l.upper() or "EJERCICIOS" in l.upper() or "TRADUCCIÓN" in l.upper() or "TŁUMACZENIE" in l.upper() or "TRANSLATION" in l.upper() or "KEY VOCABULARY" in l.upper() or "ANSWER KEY" in l.upper():
            if "SOLUCIONARIO" in l.upper() or "ANSWER KEY" in l.upper(): doc.add_page_break()
            doc.add_heading(l.replace('#', '').strip(), level=1)
            continue
            
        # Formato de Tablas
        if '|' in l and '---' not in l:
            datos = [c.strip() for l_p in l.split('|') if (c := l_p.strip())]
            if len(datos) >= 2:
                tabla = doc.add_table(rows=1, cols=2)
                tabla.style = 'Table Grid'
                cells = tabla.rows[0].cells
                cells[0].text = datos[0]
                cells[1].text = datos[1]
                continue

        p = doc.add_paragraph()
        partes = l.split('**')
        for i, parte in enumerate(partes):
            run = p.add_run(parte)
            if i % 2 == 1:
                run.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- BARRA LATERAL ---
with st.sidebar:
    st.header("🏫 Configuración")
    logo_subido = st.file_uploader("Logo de la escuela", type=["png", "jpg", "jpeg"])
    nombre_escuela = st.text_input("Escuela", "El Sabor de la Lengua")
    nombre_profe = st.text_input("Profesor/a", "Mario")
    api_key = st.text_input("API Key (Gemini)", type="password")
    st.divider()
    idioma_alumnos = st.selectbox("🌍 Idioma de los alumnos", ["Inglés", "Polaco", "Ninguno (100% Español)"])

# --- INTERFAZ ---
st.title("🎓 Generador ELE Pro")
modo = st.selectbox("Modo de generación", ["Unidad Completa (Texto + Ejercicios)", "Solo Lista de Ejercicios"])

st.divider()

col1, col2 = st.columns(2)
with col1:
    tema_input = st.text_input("Tema de la unidad")
    nivel_mcer = st.selectbox("Nivel", ["A1", "A2", "B1", "B2", "C1", "C2"])
    ext_val = "Corto"
    if modo == "Unidad Completa (Texto + Ejercicios)":
        ext_val = st.select_slider("Extensión", ["Corto", "1 pág", "2 págs", "3 págs (Extenso)"], "3 págs (Extenso)")

with col2:
    items_val = st.number_input("Items por técnica", 1, 30, 15)
    tecs_val = st.multiselect("Técnicas", ["Test de Cloze", "Preguntas de comprensión", "Verdadero o Falso", "Corregir errores", "Relacionar columnas"], default=["Relacionar columnas", "Test de Cloze"])
    gram_val = st.multiselect("Enfoque", ["Vocabulario", "Presente de Indicativo", "Pretéritos", "Futuros", "Subjuntivo", "Ser/Estar", "Por/Para", "Objetos Directos e Indirectos"], default=["Vocabulario", "Presente de Indicativo"])

# --- GENERACIÓN AUTODETECTADA ---
if st.button("🚀 Generar Material Editorial"):
    if not api_key or not tema_input:
        st.error("⚠️ Configuración incompleta.")
    else:
        try:
            with st.spinner("Paso 1: Detectando modelos autorizados para tu cuenta Pro..."):
                url_models = f"https://generativelanguage.googleapis.com/v1/models?key={api_key.strip()}"
                res_models = requests.get(url_models, timeout=20)
                
                if res_models.status_code != 200:
                    st.error(f"Error de permisos en la API: {res_models.text}")
                    st.stop()
                
                modelos_data = res_models.json().get("models", [])
                modelos_permitidos = [m["name"] for m in modelos_data if "generateContent" in m.get("supportedGenerationMethods", [])]
                
                modelo_elegido = ""
                for preferido in ["models/gemini-2.5-flash", "models/gemini-1.5-pro", "models/gemini-1.5-flash"]:
                    if preferido in modelos_permitidos:
                        modelo_elegido = preferido
                        break
                if not modelo_elegido: modelo_elegido = modelos_permitidos[0]
                
                st.info(f"✅ Conexión establecida a través de: {modelo_elegido}")
            
            with st.spinner(f"Paso 2: Redactando unidad inmersiva en {idioma_alumnos}..."):
                url_gen = f"https://generativelanguage.googleapis.com/v1/{modelo_elegido}:generateContent?key={api_key.strip()}"
                
                # INSTRUCCIONES BILINGÜES EXTREMAS
                if idioma_alumnos != "Ninguno (100% Español)":
                    instrucciones_idioma = (
                        f"REGLA DE ORO BILINGÜE: Tú eres un profesor que explica en {idioma_alumnos.upper()}.\n"
                        f"Toda la teoría, las explicaciones gramaticales, las instrucciones de los ejercicios y los títulos de las secciones DEBEN ESTAR ESCRITOS EN {idioma_alumnos.upper()}.\n"
                        f"Usa el español SOLAMENTE para:\n"
                        f"- Las listas de vocabulario (formato: Español - {idioma_alumnos}).\n"
                        f"- Los ejemplos gramaticales dentro de las explicaciones.\n"
                        f"- El contenido de los ejercicios (las frases que el alumno debe resolver).\n"
                        f"Incluye una sección dedicada a 'Traducción {idioma_alumnos.upper()}-Español' con frases del tema.\n"
                    )
                else:
                    instrucciones_idioma = "INSTRUCCIONES: Todo el material, explicaciones, títulos y ejercicios deben ser 100% en español, sin ninguna traducción.\n"

                prompt_p = (f"Actúa como profesor experto de español en la escuela {nombre_escuela}. "
                          f"Tema: {tema_input}, Nivel: {nivel_mcer}. "
                          f"{instrucciones_idioma}"
                          f"Genera un texto explicativo de {ext_val} y {items_val} ejercicios por cada técnica: {', '.join(tecs_val)}.\n"
                          f"Enfoque gramatical: {', '.join(gram_val)}.\n"
                          f"REGLA DE EJERCICIOS: Usa líneas largas '_______' para los huecos vacíos. NO escribas las respuestas dentro del ejercicio. "
                          f"Crea una sección final '# SOLUCIONARIO' (o su traducción) con todas las respuestas.\n"
                          f"Firma: {nombre_profe}.")
                
                payload = {"contents": [{"parts": [{"text": prompt_p}]}]}
                res_gen = requests.post(url_gen, json=payload, timeout=200) # Aumentado el timeout para bilingüismo complejo
                
                if res_gen.status_code == 200:
                    st.session_state['material_ia'] = res_gen.json()["candidates"][0]["content"]["parts"][0]["text"]
                    st.success("¡Material generado con éxito!")
                else:
                    st.error(f"Error {res_gen.status_code}: {res_gen.text}")

        except Exception as e:
            st.error(f"Error detectado: {e}")

if 'material_ia' in st.session_state:
    st.divider()
    docx_bytes = generar_docx_profesional(st.session_state['material_ia'], nombre_escuela, nombre_profe, tema_input, logo_file=logo_subido)
    st.download_button(f"📥 Descargar Word ({idioma_alumnos})", data=docx_bytes, file_name=f"ELE_{tema_input}.docx")
    st.markdown(st.session_state['material_ia'])
